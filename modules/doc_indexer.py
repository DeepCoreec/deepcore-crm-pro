"""doc_indexer.py — Extracción de texto e IA semántica sobre documentos.

Arquitectura dual:
  - Primario: SQLite FTS5 (siempre disponible, cero dependencias extra)
  - Secundario: LlamaIndex + Ollama (si ambos están instalados y corriendo)
"""
import os


def extraer_texto(ruta: str) -> str:
    """Extrae texto plano de PDF, TXT o DOCX."""
    ext = os.path.splitext(ruta)[1].lower()
    if ext == '.txt':
        return _leer_txt(ruta)
    elif ext == '.pdf':
        return _leer_pdf(ruta)
    elif ext in ('.docx', '.doc'):
        return _leer_docx(ruta)
    return ''


def _leer_txt(ruta: str) -> str:
    try:
        with open(ruta, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception:
        return ''


def _leer_pdf(ruta: str) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(ruta)
        return '\n'.join(p.extract_text() or '' for p in reader.pages)
    except ImportError:
        pass
    try:
        import pdfplumber
        with pdfplumber.open(ruta) as pdf:
            return '\n'.join(p.extract_text() or '' for p in pdf.pages)
    except Exception:
        pass
    return ''


def _leer_docx(ruta: str) -> str:
    try:
        import docx
        doc = docx.Document(ruta)
        return '\n'.join(p.text for p in doc.paragraphs)
    except Exception:
        return ''


def ollama_disponible() -> bool:
    """True si Ollama está corriendo en localhost."""
    try:
        import urllib.request
        urllib.request.urlopen('http://localhost:11434/api/tags', timeout=2)
        return True
    except Exception:
        return False


def buscar_con_ia(query: str, documentos: list, ollama_model: str = 'llama3.2') -> str | None:
    """
    Búsqueda semántica usando LlamaIndex + Ollama.
    documentos: [{'nombre': str, 'texto': str}]
    Retorna respuesta en español, o None si LlamaIndex/Ollama no están disponibles.
    """
    if not documentos:
        return None
    try:
        from llama_index.core import VectorStoreIndex, Document, Settings
        from llama_index.llms.ollama import Ollama
        from llama_index.embeddings.ollama import OllamaEmbedding

        Settings.llm = Ollama(
            model=ollama_model,
            base_url='http://localhost:11434',
            request_timeout=90,
        )
        Settings.embed_model = OllamaEmbedding(
            model_name=ollama_model,
            base_url='http://localhost:11434',
        )

        docs = [
            Document(text=d['texto'], metadata={'fuente': d['nombre']})
            for d in documentos
            if d.get('texto', '').strip()
        ]
        if not docs:
            return None

        index = VectorStoreIndex.from_documents(docs)
        engine = index.as_query_engine(similarity_top_k=3)
        respuesta = engine.query(
            f"Responde en español: {query}\nCita el documento fuente cuando sea relevante."
        )
        return str(respuesta).strip()
    except ImportError:
        return None
    except Exception:
        return None
